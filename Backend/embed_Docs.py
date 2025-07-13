import os
import io
import json
import pytesseract
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LCDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account
from pathlib import Path

# ------------------- CONFIG -------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, "temp_downloads")
EMBEDDING_FOLDER = os.path.join(BASE_DIR, "embeddings")
PROCESSED_FILES_JSON = os.path.join(EMBEDDING_FOLDER, "processed_files.json")

FOLDER_ID = "1iCLFeSEuYCsi0gf3eUFo4Y3P8PLyMrT7"  # Your Drive folder ID
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"

SERVICE_ACCOUNT_FILE = Path(__file__).resolve().parent / "Config" / "file-service.json"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(EMBEDDING_FOLDER, exist_ok=True)

# ------------------- GOOGLE DRIVE -------------------
def init_drive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    return build('drive', 'v3', credentials=creds)

def list_files_in_folder(drive_service, folder_id):
    files = []
    page_token = None
    while True:
        response = drive_service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="nextPageToken, files(id, name, mimeType, modifiedTime)",
            pageToken=page_token
        ).execute()
        files.extend(response.get('files', []))
        page_token = response.get('nextPageToken', None)
        if page_token is None:
            break
    return files

def download_file(drive_service, file_id, filename):
    request = drive_service.files().get_media(fileId=file_id)
    filepath = os.path.join(TEMP_DIR, filename)
    with open(filepath, "wb") as f:
        downloader = MediaIoBaseDownload(f, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
    return filepath

# ------------------- TEXT EXTRACTION -------------------
def extract_text_from_pdf(path):
    text = ""
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            text += page.extract_text() or ""
    except:
        pass
    if text.strip():
        return text
    try:
        images = convert_from_path(path)
        for image in images:
            text += pytesseract.image_to_string(image)
    except Exception as e:
        print(f"OCR failed for {path}: {e}")
    return text

def extract_text_from_docx(path):
    try:
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except:
        return ""

def extract_text_from_xlsx(path):
    try:
        df = pd.read_excel(path, sheet_name=None)
        all_text = ""
        for sheet in df.values():
            all_text += sheet.astype(str).apply(lambda x: ' '.join(x), axis=1).str.cat(sep='\n')
        return all_text
    except:
        return ""

# ------------------- UTILS -------------------
def load_processed_files_metadata():
    if os.path.exists(PROCESSED_FILES_JSON):
        with open(PROCESSED_FILES_JSON, "r") as f:
            return json.load(f)
    return {}

def save_processed_files_metadata(metadata):
    with open(PROCESSED_FILES_JSON, "w") as f:
        json.dump(metadata, f, indent=2)

# ------------------- EMBEDDING PIPELINE -------------------

    drive_service = init_drive_service()
    files = list_files_in_folder(drive_service, FOLDER_ID)

    if not files:
        print("No files found in the Drive folder.")
        return

    print(f"\n📥 Found {len(files)} files in Drive folder.")

    processed_files = load_processed_files_metadata()

    # Filter files that are new or modified
    files_to_process = []
    for f in files:
        filename = f["name"]
        modified_time = f["modifiedTime"]
        # Compare stored modified time to detect changes
        if filename not in processed_files or processed_files[filename] != modified_time:
            files_to_process.append(f)

    if not files_to_process:
        print("No new or modified files to process. Loading existing FAISS index...")
        if os.path.exists(os.path.join(EMBEDDING_FOLDER, "index.faiss")):
            # Load existing FAISS vectorstore
            embedding_model = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
            vectorstore = FAISS.load_local(EMBEDDING_FOLDER, embedding_model)
            print("✅ Loaded existing FAISS vectorstore.")
        else:
            print("❌ No existing vectorstore found. Please add or modify files to build index.")
        return

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,separators=["\n\n", "\n", ".", " ", ""])
    embedding_model = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    all_docs = []

    for file in files_to_process:
        filename = file["name"]
        file_id = file["id"]
        mime = file["mimeType"]
        modified_time = file["modifiedTime"]

        if not (filename.endswith(".pdf") or filename.endswith(".docx") or filename.endswith(".xlsx")):
            print(f"⛔ Skipping unsupported file type: {filename}")
            continue

        print(f"⬇️ Downloading: {filename}")
        path = download_file(drive_service, file_id, filename)

        # Extract text
        if filename.endswith(".pdf"):
            raw_text = extract_text_from_pdf(path)
        elif filename.endswith(".docx"):
            raw_text = extract_text_from_docx(path)
        elif filename.endswith(".xlsx"):
            raw_text = extract_text_from_xlsx(path)
        else:
            raw_text = ""

        if not raw_text.strip():
            print(f"⚠️ Skipping empty/unreadable file: {filename}")
            continue

        chunks = splitter.split_text(raw_text)
        for chunk in chunks:
            all_docs.append(LCDocument(page_content=chunk, metadata={"source": filename}))

        # Update processed_files metadata right after processing each file
        processed_files[filename] = modified_time

    if not all_docs:
        print("No new chunks to add to vectorstore.")
        return

    # Load existing vectorstore if exists, else create new
    if os.path.exists(os.path.join(EMBEDDING_FOLDER, "index.faiss")):
        print("Loading existing FAISS vectorstore to add new docs...")
        vectorstore = FAISS.load_local(EMBEDDING_FOLDER, embedding_model)
        vectorstore.add_documents(all_docs)
    else:
        print("Creating new FAISS vectorstore...")
        vectorstore = FAISS.from_documents(all_docs, embedding_model)

    vectorstore.save_local(EMBEDDING_FOLDER)
    save_processed_files_metadata(processed_files)

    print(f"\n✅ FAISS vectorstore saved/updated to `{EMBEDDING_FOLDER}`")
    print(f"Processed files metadata updated.")

# ------------------- RUN -------------------
