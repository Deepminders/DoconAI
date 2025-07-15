from bson import ObjectId
from fastapi import File, UploadFile, HTTPException, Form
from typing import Optional, List
from fastapi.responses import StreamingResponse, RedirectResponse, JSONResponse
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from google.oauth2 import service_account
from Config.db import document_collection, get_next_sequence_value, staff_collection, user_collection
from pathlib import Path
from datetime import datetime
import httpx
import fitz  # PyMuPDF
import io
import requests
import os
import tempfile
import logging
import docx # Required for reading .docx files
import google.generativeai as genai
from dotenv import load_dotenv
import time
import re
import pandas as pd
import openpyxl
from io import StringIO

# --- Configuration ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- LLM and API Key Configuration ---
try:
    load_dotenv()
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        logger.warning("GOOGLE_API_KEY not found in .env file. LLM classification will be disabled.")
        genai.configure(api_key="DUMMY_KEY") # Configure with a dummy to avoid errors on import
    else:
        genai.configure(api_key=GOOGLE_API_KEY)
except ImportError:
    logger.warning("'python-dotenv' not installed. Cannot load GOOGLE_API_KEY from .env file.")
    GOOGLE_API_KEY = None


# --- Google Drive Configuration ---
DRIVE_ROOT_FOLDER_ID = "1GxNapTLGFcUmshr3ZEjBHSVy3BW8NdJr" 
CURRENT_USER = "Sample User"

# --- Service Account Initialization ---
SERVICE_ACCOUNT_FILE = Path(__file__).resolve().parent / '..' / "Config" / "file-service.json"
SCOPES = ['https://www.googleapis.com/auth/drive.file']
try:
    credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    drive_service = build('drive', 'v3', credentials=credentials)
    logger.info("Google Drive service initialized successfully.")
except Exception as e:
    logger.error(f"Failed to initialize Google Drive service: {e}")
    drive_service = None

# --- Keyword-based Classifier (Fallback) ---
def classify_document_simple(filename: str) -> str:
    """A simple keyword-based classifier used as a fallback."""
    filename_lower = filename.lower()
    if "boq" in filename_lower or "bill of quantities" in filename_lower: return "Bill of Quantities (BOQ)"
    if "contract" in filename_lower or "agreement" in filename_lower: return "Contracts and Agreements"
    if "tender" in filename_lower or "bid" in filename_lower: return "Tender Documents"
    if "progress" in filename_lower and "report" in filename_lower: return "Progress Reports"
    if "final" in filename_lower and "report" in filename_lower: return "Final Reports"
    if "cost" in filename_lower or "estimation" in filename_lower: return "Cost Estimations"
    if "invoice" in filename_lower or "payment" in filename_lower: return "Invoices and Financials"
    if "drawing" in filename_lower or "plan" in filename_lower: return "Drawings and Plans"
    if "permit" in filename_lower or "license" in filename_lower: return "Permits and Licenses"
    if "safety" in filename_lower or "compliance" in filename_lower: return "Safety and Compliance"
    return "Other"

# --- LLM-Powered Document Classifier ---
async def classify_document_with_llm(filename: str, content: bytes, file_type: str) -> str:
    """
    Predicts document category using the Gemini LLM with a few-shot prompt.
    Falls back to simple keyword-based classification if the API call fails or is not configured.
    """
    if not GOOGLE_API_KEY:
        logger.info("Using simple classifier because GOOGLE_API_KEY is not configured.")
        return classify_document_simple(filename)

    text_content = ""
    try:
        if file_type == "application/pdf":
            with fitz.open(stream=io.BytesIO(content), filetype="pdf") as doc:
                text_content = "".join(page.get_text() for page in doc)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx'):
            doc = docx.Document(io.BytesIO(content))
            text_content = "\n".join([para.text for para in doc.paragraphs])
        else:
            # Attempt to decode as text for other file types (e.g., .txt, .csv)
            text_content = content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}. Using fallback classifier.")
        return classify_document_simple(filename)

    if not text_content.strip():
        logger.warning(f"No text content could be extracted from {filename}. Using fallback classifier.")
        return classify_document_simple(filename)

    # Truncate content to a reasonable length for the API
    truncated_content = text_content[:8000]

    categories = [
        "Bill of Quantities (BOQ)", "Contracts and Agreements", "Tender Documents",
        "Progress Reports", "Final Reports", "Cost Estimations", "Invoices and Financials",
        "Drawings and Plans", "Permits and Licenses", "Safety and Compliance", "Other"
    ]
    
    # Few-shot prompt with high-quality examples to improve accuracy
    prompt = f"""
    You are an expert document classifier for the construction industry. Your task is to analyze the provided document content and classify it into exactly one of the following categories. Respond with ONLY the category name and nothing else.

    Categories: {', '.join(categories)}

    ---
    **Example 1:**
    Document Content: "This Agreement is made and entered into as of July 10, 2025, by and between Apex Construction Ltd. ('Contractor') and Elysian Properties Inc. ('Owner'). The Contractor shall provide all labor, materials, and equipment for the construction of the 'Azure Tower' project..."
    Category: Contracts and Agreements

    --- 
    **Example 2:**
    Document Content: "Weekly Progress Report #12. Period: July 1 - July 7, 2025. Project: Azure Tower. Summary: Concrete pouring for the 5th floor is 80% complete. Electrical rough-in is ongoing on floors 2 and 3. Safety inspection conducted on July 5, with no major findings. Upcoming Activities: Complete 5th-floor slab, begin framing on the 6th floor."
    Category: Progress Reports

    ---
    **Example 3:**
    Document Content: "Bill of Quantities - Item No. 3.1. Description: Reinforced Concrete Grade C35. Unit: Cubic Meter. Quantity: 1,200. Rate (USD): 150.00. Total (USD): 180,000.00."
    Category: Bill of Quantities (BOQ)

    ---
    **Document to Classify:**
    Document Filename: "{filename}"
    Document Content (excerpt):
    {truncated_content}
    ---
    Category:
    """

    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = await model.generate_content_async(prompt)
        predicted_category = response.text.strip()

        # Validate that the model's response is one of the allowed categories
        if predicted_category in categories:
            logger.info(f"LLM classified '{filename}' as '{predicted_category}'")
            return predicted_category
        else:
            logger.warning(f"LLM returned an invalid category: '{predicted_category}'. Using fallback.")
            return classify_document_simple(filename)
    except Exception as e:
        logger.error(f"Gemini API call failed: {e}. Using fallback classifier.")
        return classify_document_simple(filename)


# --- Google Drive Helper Functions ---
def get_folder_id_by_name(folder_name: str, parent_id: str) -> str:
    """Get folder ID by name within a parent folder."""
    try:
        query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and '{parent_id}' in parents and trashed=false"
        response = drive_service.files().list(q=query, fields="files(id)").execute()
        files = response.get('files', [])
        return files[0]['id'] if files else None
    except Exception as e:
        logger.error(f"Error getting folder ID for {folder_name}: {e}")
        return None


def is_folder_empty(folder_id: str) -> bool:
    """Check if a folder is empty (contains no files or folders)."""
    try:
        query = f"'{folder_id}' in parents and trashed=false"
        response = drive_service.files().list(q=query, fields="files(id)", pageSize=1).execute()
        files = response.get('files', [])
        return len(files) == 0
    except Exception as e:
        logger.error(f"Error checking if folder {folder_id} is empty: {e}")
        return False


def get_or_create_drive_folder(name: str, parent_id: str) -> str:
    """Checks for a folder in Google Drive by name within a parent. Creates it if it doesn't exist."""
    if not drive_service:
        raise HTTPException(status_code=503, detail="Google Drive service is not available")
    
    try:
        # First, try to find existing folder
        existing_folder_id = get_folder_id_by_name(name, parent_id)
        if existing_folder_id:
            return existing_folder_id
        
        # Create new folder if it doesn't exist
        folder_metadata = {
            'name': name, 
            'mimeType': 'application/vnd.google-apps.folder', 
            'parents': [parent_id]
        }
        
        # Add retry mechanism for folder creation
        max_retries = 3
        for attempt in range(max_retries):
            try:
                folder = drive_service.files().create(body=folder_metadata, fields='id').execute()
                folder_id = folder.get('id')
                logger.info(f"Created new folder '{name}' with ID: {folder_id}")
                return folder_id
            except Exception as e:
                if "SSL" in str(e) or "EOF" in str(e) or "Connection" in str(e):
                    if attempt < max_retries - 1:
                        logger.warning(f"Network error creating folder on attempt {attempt + 1}/{max_retries}: {e}. Retrying...")
                        time.sleep(2 ** attempt)  # Exponential backoff
                        continue
                    else:
                        raise HTTPException(
                            status_code=503, 
                            detail="Network connectivity issue with Google Drive. Please check your internet connection and try again."
                        )
                else:
                    raise e
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Google Drive API error: {e}")
        if "SSL" in str(e) or "EOF" in str(e):
            raise HTTPException(
                status_code=503, 
                detail="Network connectivity issue with Google Drive. Please check your internet connection and try again."
            )
        elif "credentials" in str(e).lower():
            raise HTTPException(
                status_code=503, 
                detail="Google Drive authentication failed. Please check service account credentials."
            )
        else:
            raise HTTPException(
                status_code=503, 
                detail=f"Google Drive service error: {str(e)}"
            )


async def cleanup_empty_folders(document):
    """Clean up empty folders after deleting document files."""
    try:
        project_id = document.get("project_id")
        category = document.get("document_category")
        
        if not project_id or not category:
            return
        
        # Get project folder ID
        project_folder_id = get_folder_id_by_name(project_id, DRIVE_ROOT_FOLDER_ID)
        if not project_folder_id:
            return
        
        # Get category folder ID
        category_folder_id = get_folder_id_by_name(category, project_folder_id)
        if not category_folder_id:
            return
        
        # Check if category folder is empty
        if is_folder_empty(category_folder_id):
            try:
                drive_service.files().delete(fileId=category_folder_id).execute()
                logger.info(f"Deleted empty category folder: {category}")
            except Exception as e:
                logger.warning(f"Could not delete empty category folder {category}: {e}")
        
        # Check if project folder is empty (after potentially deleting category folder)
        if is_folder_empty(project_folder_id):
            try:
                drive_service.files().delete(fileId=project_folder_id).execute()
                logger.info(f"Deleted empty project folder: {project_id}")
            except Exception as e:
                logger.warning(f"Could not delete empty project folder {project_id}: {e}")
                
    except Exception as e:
        logger.warning(f"Error during folder cleanup: {e}")


# --- API Endpoints ---

async def classify_and_stage_document(file: UploadFile):
    """
    Step 1 of the upload process. Classifies the document and saves it temporarily.
    """
    content = await file.read()
    predicted_category = await classify_document_with_llm(file.filename, content, file.content_type)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as tmp:
        tmp.write(content)
        temp_file_path = tmp.name
        
    return JSONResponse({
        "status": "pending_confirmation",
        "message": "Document classification predicted. Please confirm.",
        "predicted_category": predicted_category,
        "original_filename": file.filename,
        "temp_file_path": temp_file_path
    })

async def addDocument(
    proj_id: str,
    doc_name: str,
    confirmed_category: str,
    temp_file_path: str,
    original_filename: str,
    user_id: str  # Added user_id parameter
):
    """
    Step 2 of the upload process. Confirms the category and saves the document.
    Now includes user_id for tracking who uploaded the document.
    """
    if not os.path.exists(temp_file_path):
        raise HTTPException(status_code=404, detail="Temporary file not found. The session may have expired. Please upload again.")

    with open(temp_file_path, "rb") as f:
        content = f.read()

    file_content_type = "application/octet-stream"
    if original_filename.lower().endswith('.pdf'): file_content_type = "application/pdf"
    elif original_filename.lower().endswith('.docx'): file_content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif original_filename.lower().endswith('.xls'): file_content_type = "application/vnd.ms-excel"
    elif original_filename.lower().endswith('.xlsx'): file_content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    page_count = 0
    if file_content_type == "application/pdf":
        try:
            with fitz.open(stream=io.BytesIO(content), filetype="pdf") as doc:
                page_count = doc.page_count
        except Exception as e:
            logger.warning(f"Could not get page count for PDF: {e}")

    project_folder_id = get_or_create_drive_folder(proj_id, DRIVE_ROOT_FOLDER_ID)
    category_folder_id = get_or_create_drive_folder(confirmed_category, project_folder_id)

    file_stream = io.BytesIO(content)
    media = MediaIoBaseUpload(file_stream, mimetype=file_content_type, resumable=True)
    
    drive_response = drive_service.files().create(
        body={"name": original_filename, "parents": [category_folder_id]},
        media_body=media,
        fields="id, size, modifiedTime, webViewLink"
    ).execute()
    
    file_id = drive_response.get("id")
    document_link = drive_response.get("webViewLink")
    download_link = f"https://drive.google.com/uc?export=download&id={file_id}"

    document_id = await get_next_sequence_value("document_id")
    current_time = datetime.now().isoformat()
    
    version_info = {
        "version": 1, 
        "google_drive_id": file_id, 
        "original_filename": original_filename,
        "document_size": int(drive_response.get("size", len(content))), 
        "upload_date": current_time,
        "last_modified_date": drive_response.get("modifiedTime", current_time), 
        "page_count": page_count,
        "file_type": file_content_type, 
        "uploaded_by": user_id,  # Store user_id instead of CURRENT_USER
        "document_link": document_link, 
        "download_link": download_link
    }
    
    file_info = {
        "document_id": document_id, 
        "project_id": proj_id, 
        "document_name": doc_name,
        "document_category": confirmed_category, 
        "current_version": 1, 
        "last_modified_date": current_time,
        "document_link": document_link, 
        "download_link": download_link, 
        "created_by": user_id,  # Add who created the document
        "versions": [version_info]
    }

    try:
        await document_collection.insert_one(file_info)
        os.remove(temp_file_path)
        return JSONResponse({
            "status": "success", 
            "message": "Document added successfully",
            "document_id": file_info["document_id"], 
            "drive_id": file_id, 
            "version": 1,
            "created_by": user_id,  # Return user_id in response
            "links": {"view": document_link, "download": download_link}
        })
    except Exception as e:
        drive_service.files().delete(fileId=file_id).execute()
        os.remove(temp_file_path)
        raise HTTPException(status_code=500, detail=f"Database insertion failed: {str(e)}")


async def replaceDocument(docid: str, file: UploadFile, new_name: str = None, confirmed_category: str = None, user_id: str = None):
    """
    Replaces an existing document, creating a new version in Drive and the database.
    Now includes user_id to track who made the replacement.
    """
    try:
        document_id = int(docid)
        existing_doc = await document_collection.find_one({"document_id": document_id})
        if not existing_doc:
            raise HTTPException(status_code=404, detail="Document not found")

        content = await file.read()
        
        # Get page count for PDF files
        page_count = 0
        if file.content_type == "application/pdf":
            try:
                with fitz.open(stream=io.BytesIO(content), filetype="pdf") as doc:
                    page_count = doc.page_count
            except Exception as e:
                logger.warning(f"Could not get page count for PDF: {e}")

        # Get project folder and create/get category folder
        project_folder_id = get_or_create_drive_folder(existing_doc["project_id"], DRIVE_ROOT_FOLDER_ID)
        category_folder_id = get_or_create_drive_folder(confirmed_category, project_folder_id)

        # Create a versioned filename
        new_version_number = existing_doc["current_version"] + 1
        base_name, extension = os.path.splitext(file.filename)
        new_filename = f"{base_name}_v{new_version_number}{extension}"

        # Upload the new file to Google Drive with retry mechanism
        file_stream = io.BytesIO(content)
        media = MediaIoBaseUpload(file_stream, mimetype=file.content_type, resumable=True)
        
        max_retries = 3
        drive_response = None
        
        for attempt in range(max_retries):
            try:
                drive_response = drive_service.files().create(
                    body={"name": new_filename, "parents": [category_folder_id]},
                    media_body=media,
                    fields="id, size, modifiedTime, webViewLink"
                ).execute()
                break
            except Exception as e:
                if "SSL" in str(e) or "EOF" in str(e) or "Connection" in str(e):
                    if attempt < max_retries - 1:
                        logger.warning(f"Network error uploading file on attempt {attempt + 1}/{max_retries}: {e}. Retrying...")
                        # Reset the file stream for retry
                        file_stream = io.BytesIO(content)
                        media = MediaIoBaseUpload(file_stream, mimetype=file.content_type, resumable=True)
                        time.sleep(2 ** attempt)  # Exponential backoff
                        continue
                    else:
                        raise HTTPException(
                            status_code=503, 
                            detail="Network connectivity issue with Google Drive. Please check your internet connection and try again."
                        )
                else:
                    raise HTTPException(
                        status_code=503, 
                        detail=f"Failed to upload file to Google Drive: {str(e)}"
                    )
        
        if not drive_response:
            raise HTTPException(status_code=503, detail="Failed to upload file to Google Drive after multiple attempts")
        
        file_id = drive_response.get("id")
        document_link = drive_response.get("webViewLink")
        download_link = f"https://drive.google.com/uc?export=download&id={file_id}"

        # Prepare the new version's metadata for MongoDB
        current_time = datetime.now().isoformat()
        new_version_data = {
            "version": new_version_number, 
            "google_drive_id": file_id, 
            "original_filename": file.filename,
            "document_size": int(drive_response.get("size", len(content))), 
            "upload_date": current_time,
            "last_modified_date": drive_response.get("modifiedTime", current_time), 
            "page_count": page_count,
            "file_type": file.content_type, 
            "uploaded_by": user_id if user_id else "Unknown User",  # Store user_id
            "document_link": document_link, 
            "download_link": download_link
        }

        # Update the document in MongoDB with all fields that need updating
        update_query = {
            "$set": {
                "current_version": new_version_number,
                "last_modified_date": new_version_data["last_modified_date"],
                "document_link": document_link, 
                "download_link": download_link,
                "document_category": confirmed_category,  # Update category
                "last_modified_by": user_id if user_id else "Unknown User"  # Track who modified
            },
            "$push": {"versions": new_version_data}
        }
        
        # Update document name if provided
        if new_name:
            update_query["$set"]["document_name"] = new_name

        result = await document_collection.update_one({"document_id": document_id}, update_query)
        
        if result.matched_count == 0:
            # If update failed, clean up the uploaded file
            try:
                drive_service.files().delete(fileId=file_id).execute()
            except Exception as cleanup_error:
                logger.error(f"Failed to cleanup uploaded file after database update failure: {cleanup_error}")
            raise HTTPException(status_code=404, detail="Document not found during update")

        # Get the updated document to return fresh data
        updated_doc = await document_collection.find_one({"document_id": document_id}, {"_id": 0})

        return JSONResponse({
            "status": "success", 
            "message": f"Document replaced successfully. New version is v{new_version_number}",
            "document_id": document_id, 
            "new_version": new_version_number, 
            "drive_id": file_id,
            "modified_by": user_id,  # Return who modified it
            "updated_document": updated_doc,  # Return full updated document for frontend refresh
            "updated_fields": {
                "document_name": new_name if new_name else existing_doc["document_name"],
                "document_category": confirmed_category,
                "version": new_version_number,
                "document_link": document_link,
                "download_link": download_link,
                "last_modified_date": new_version_data["last_modified_date"],
                "last_modified_by": user_id
            }
        })

    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Unexpected error in replaceDocument: {e}")
        raise HTTPException(
            status_code=500, 
            detail=f"An unexpected error occurred while replacing the document: {str(e)}"
        )
        
async def downloadDocument(docid: str, version: Optional[int] = None):
    """Download a document by streaming it through the API."""
    
    doc = await document_collection.find_one({"document_id": int(docid)})
    if not doc: 
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Find the correct version
    version_to_download = (doc["versions"][-1] if version is None 
                          else next((v for v in doc["versions"] if v["version"] == version), None))
    
    if not version_to_download or "download_link" not in version_to_download:
        raise HTTPException(status_code=404, detail=f"Download link for version {version or 'latest'} not found.")
    
    # Stream the file with redirect following
    async with httpx.AsyncClient(follow_redirects=True) as client:  # Enable redirects
        try:
            response = await client.get(version_to_download["download_link"])
            response.raise_for_status()
            
            # Get filename from document or create default
            filename = version_to_download.get("filename", f"document_{docid}.pdf")
            
            return StreamingResponse(
                io.BytesIO(response.content),
                media_type="application/octet-stream",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        except httpx.HTTPStatusError as e:
            raise HTTPException(status_code=500, detail=f"Failed to download file: {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Download error: {str(e)}")
        
async def viewDocument(docid: str, version: Optional[int] = None):
    """Redirects to the viewable link for a specific document version."""
    doc = await document_collection.find_one({"document_id": int(docid)})
    if not doc: raise HTTPException(status_code=404, detail="Document not found")

    # Find the correct version, defaulting to the latest
    version_to_view = doc["versions"][-1] if version is None else next((v for v in doc["versions"] if v["version"] == version), None)
    
    if not version_to_view or "document_link" not in version_to_view:
        raise HTTPException(status_code=404, detail=f"View link for version {version or 'latest'} not found.")
        
    return RedirectResponse(url=version_to_view["document_link"])


async def updateDocument(docid: str, new_name: Optional[str] = None, new_category: Optional[str] = None):
    """Updates a document's metadata, such as its name or category."""
    if new_name is None and new_category is None: 
        raise HTTPException(status_code=400, detail="No update data provided")
    
    update_data = {"last_modified_date": datetime.now().isoformat()}
    if new_name is not None: update_data["document_name"] = new_name
    if new_category is not None: update_data["document_category"] = new_category
    
    result = await document_collection.update_one({"document_id": int(docid)}, {"$set": update_data})
    if result.matched_count == 0: 
        raise HTTPException(status_code=404, detail="Document not found")
    
    updated_doc = await document_collection.find_one({"document_id": int(docid)}, {"_id": 0})
    return JSONResponse({
        "status": "success", 
        "message": "Document updated successfully", 
        "document": updated_doc,
        "updated_fields": {k: v for k, v in update_data.items() if k != "last_modified_date"}
    })

async def deleteDocument(docid: str):
    """Deletes a document and all its versions from Google Drive and MongoDB."""
    document = await document_collection.find_one({"document_id": int(docid)})
    if not document: 
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not drive_service:
        raise HTTPException(status_code=503, detail="Google Drive service is not available")
    
    deleted_files = []
    failed_deletions = []
    
    # Delete all file versions from Google Drive
    for version in document.get("versions", []):
        file_id = version.get("google_drive_id")
        if not file_id:
            continue
            
        try:
            # Add retry mechanism for network issues
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    drive_service.files().delete(fileId=file_id).execute()
                    deleted_files.append(file_id)
                    logger.info(f"Successfully deleted file {file_id} from Google Drive")
                    break
                except Exception as e:
                    if "SSL" in str(e) or "EOF" in str(e) or "Connection" in str(e):
                        if attempt < max_retries - 1:
                            logger.warning(f"Network error on attempt {attempt + 1}/{max_retries} for file {file_id}: {e}. Retrying...")
                            time.sleep(2 ** attempt)  # Exponential backoff
                            continue
                        else:
                            logger.error(f"Failed to delete file {file_id} after {max_retries} attempts due to network issues: {e}")
                            failed_deletions.append({"file_id": file_id, "error": str(e)})
                    else:
                        logger.error(f"Could not delete file {file_id} from Drive: {e}")
                        failed_deletions.append({"file_id": file_id, "error": str(e)})
                    break
        except Exception as e:
            logger.error(f"Unexpected error deleting file {file_id}: {e}")
            failed_deletions.append({"file_id": file_id, "error": str(e)})
    
    # Now check and clean up empty folders
    try:
        await cleanup_empty_folders(document)
    except Exception as e:
        logger.warning(f"Failed to cleanup empty folders: {e}")
    
    # Delete the document record from MongoDB
    await document_collection.delete_one({"document_id": int(docid)})
    
    # Prepare response
    response_data = {
        "status": "success", 
        "message": "Document deleted successfully",
        "deleted_files_count": len(deleted_files),
        "failed_deletions_count": len(failed_deletions)
    }
    
    if failed_deletions:
        response_data["warning"] = f"Document deleted from database, but {len(failed_deletions)} file(s) could not be deleted from Google Drive due to network issues"
        response_data["failed_files"] = failed_deletions
    
    return JSONResponse(response_data)

async def fetchRecents():
    """Fetches the 10 most recently modified documents."""
    cursor = document_collection.find({}, {"_id": 0}).sort("last_modified_date", -1).limit(10)
    recent_docs = await cursor.to_list(length=10)
    return JSONResponse({"status": "success", "recent_documents": recent_docs})

async def getDocInfo(docid: int):
    """Retrieves all information for a single document by its ID."""
    doc = await document_collection.find_one({"document_id": docid}, {"_id": 0})
    if not doc: raise HTTPException(status_code=404, detail="Document not found")
    return JSONResponse({"status": "success", "document": doc})

async def getDocsfromProject(proj_id: str):
    """Retrieves all documents associated with a specific project ID."""
    cursor = document_collection.find({"project_id": proj_id}, {"_id": 0})
    project_docs = await cursor.to_list(length=None)
    return JSONResponse({"status": "success", "documents": project_docs})



#---- laavanjan's document route for direct download ----

async def proxy_download_document(docid: str, version: Optional[int] = None):
    doc = await document_collection.find_one({"document_id": int(docid)})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Find the correct version, defaulting to the latest
    version_to_download = doc["versions"][-1] if version is None else next(
        (v for v in doc["versions"] if v["version"] == version), None
    )

    if not version_to_download or "download_link" not in version_to_download:
        raise HTTPException(status_code=404, detail=f"Download link for version {version or 'latest'} not found.")

    # Proxy the file from Google Drive
    download_url = version_to_download["download_link"]
    r = requests.get(download_url, stream=True)
    if r.status_code != 200:
        raise HTTPException(status_code=502, detail="Failed to fetch file from Google Drive.")

    return StreamingResponse(
        r.iter_content(chunk_size=8192),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc["document_name"]}.pdf"'}
    )

async def fetchUserDocuments(user_id: str):
    """Fetches ALL documents uploaded by a specific user."""
    try:
        cursor = document_collection.find({"created_by": user_id}, {"_id": 0}).sort("last_modified_date", -1)
        user_docs = await cursor.to_list(length=None)
        
        logger.info(f"Found {len(user_docs)} documents for user {user_id}")
        return JSONResponse({
            "status": "success", 
            "recent_documents": user_docs,
            "user_id": user_id,
            "count": len(user_docs)
        })
    except Exception as e:
        logger.error(f"Error fetching documents for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch user documents: {str(e)}")
    
async def fetchProjectDocumentsByUser(user_id: str):
    """
    Fetches ALL documents from projects assigned to a user if they are Project Manager or Project Owner.
    For other roles, returns only their own documents.
    """
    try:
        # First, find the user in staff collection
        staff_member = await staff_collection.find_one({"_id": ObjectId(user_id)})
        
        if not staff_member:
            # raise HTTPException(status_code=404, detail="Staff member not found")
            staff_member = await user_collection.find_one({"_id": ObjectId(user_id)})
        
        user_role = staff_member.get("staff_role", "").strip() or staff_member.get("user_role", "").strip()
        assigned_projects = staff_member.get("assigned_projects", []) or  []
        
        logger.info(f"User {user_id} has role: {user_role} and assigned projects: {assigned_projects}")
        
        # Check if user is Project Manager or Project Owner
        if user_role in ["Project Manager", "Project owner"]:
            # Get all project IDs as strings (convert ObjectId to string for comparison)
            project_ids = [str(project_id) for project_id in assigned_projects]
            
            if not project_ids:
                logger.info(f"No projects assigned to user {user_id}")
                return JSONResponse({
                    "status": "success",
                    "recent_documents": [],
                    "user_id": user_id,
                    "user_role": user_role,
                    "assigned_projects": [],
                    "count": 0,
                    "message": "No projects assigned to this user"
                })
            
            # Fetch all documents from assigned projects
            cursor = document_collection.find(
                {"project_id": {"$in": project_ids}}, 
                {"_id": 0}
            ).sort("last_modified_date", -1)
            
            project_docs = await cursor.to_list(length=None)
            
            logger.info(f"Found {len(project_docs)} documents across {len(project_ids)} projects for {user_role} {user_id}")
            
            return JSONResponse({
                "status": "success",
                "recent_documents": project_docs,
                "user_id": user_id,
                "user_role": user_role,
                "assigned_projects": project_ids,
                "count": len(project_docs),
                "access_level": "project_manager"
            })
        
        else:
            # For other roles, return only their own documents
            logger.info(f"User {user_id} has role {user_role}, returning only own documents")
            
            cursor = document_collection.find(
                {"created_by": user_id}, 
                {"_id": 0}
            ).sort("last_modified_date", -1)
            
            user_docs = await cursor.to_list(length=None)
            
            return JSONResponse({
                "status": "success",
                "recent_documents": user_docs,
                "user_id": user_id,
                "user_role": user_role,
                "assigned_projects": [],
                "count": len(user_docs),
                "access_level": "own_documents_only"
            })
            
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Error fetching project documents for user {user_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch project documents: {str(e)}")

async def summarizeDocument(doc_id: str):
    """
    Fetches a document by ID, reads its content, and generates a clean, formatted AI summary.
    Now supports Excel files (BOQ, cost sheets, etc.) with enhanced financial analysis.
    """
    try:
        # Fetch document from database
        document_record = await document_collection.find_one({"document_id": int(doc_id)})
        if not document_record:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get the latest version's download link
        latest_version = document_record["versions"][-1] if document_record.get("versions") else None
        if not latest_version or "download_link" not in latest_version:
            raise HTTPException(status_code=404, detail="Document download link not found")
        
        download_url = latest_version["download_link"]
        original_filename = latest_version.get("original_filename", "document")
        file_type = latest_version.get("file_type", "")
        
        logger.info(f"Downloading document from: {download_url}")
        
        # Download the file content
        response = requests.get(download_url, stream=True)
        if response.status_code != 200:
            raise HTTPException(status_code=502, detail="Failed to download document from Google Drive")
        
        # Extract text content based on file type
        text_content = ""
        
        if file_type == "application/pdf" or original_filename.lower().endswith('.pdf'):
            # Extract text from PDF using PyMuPDF (fitz)
            try:
                pdf_bytes = io.BytesIO(response.content)
                with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_doc:
                    for page_num in range(pdf_doc.page_count):
                        page = pdf_doc[page_num]
                        text_content += page.get_text() + "\n"
                        
                logger.info(f"Extracted {len(text_content)} characters from PDF using PyMuPDF")
                
            except Exception as e:
                logger.error(f"Failed to extract PDF content with PyMuPDF: {e}")
                raise HTTPException(status_code=500, detail="Failed to extract PDF content")
                
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or original_filename.lower().endswith('.docx'):
            # Extract text from DOCX
            try:
                docx_file = io.BytesIO(response.content)
                docx_doc = docx.Document(docx_file)
                
                for paragraph in docx_doc.paragraphs:
                    text_content += paragraph.text + "\n"
                    
                # Also extract text from tables
                for table in docx_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + "\t"
                        text_content += "\n"
                        
                logger.info(f"Extracted {len(text_content)} characters from DOCX")
                
            except Exception as e:
                logger.error(f"Failed to extract DOCX content: {e}")
                raise HTTPException(status_code=500, detail="Failed to extract DOCX content")
                
        elif file_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"] or original_filename.lower().endswith(('.xls', '.xlsx')):
            # Enhanced BOQ/Excel analysis - FIXED VERSION
            try:
                excel_file = io.BytesIO(response.content)
                
                # Read Excel file with pandas
                if original_filename.lower().endswith('.xlsx'):
                    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
                else:
                    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                
                # Initialize analysis variables
                grand_total_budget = 0
                all_cost_items = []
                sheet_summaries = []
                
                # Process all sheets
                for sheet_name, df in excel_data.items():
                    text_content += f"\n=== SHEET: {sheet_name} ===\n"
                    
                    if not df.empty:
                        # Basic sheet info
                        text_content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                        text_content += "COLUMNS: " + ", ".join(str(col) for col in df.columns) + "\n\n"
                        
                        # Find potential cost/amount columns with more flexible matching
                        cost_columns = []
                        quantity_columns = []
                        item_columns = []
                        
                        for col in df.columns:
                            col_str = str(col).lower().strip()
                            # Look for cost/amount related columns (more flexible)
                            if any(keyword in col_str for keyword in ['amount', 'cost', 'price', 'total', 'value', 'budget', 'sum', 'rate', 'rs', 'rupee', 'dollar', 'usd', 'lkr']):
                                cost_columns.append(col)
                            # Look for quantity columns
                            elif any(keyword in col_str for keyword in ['qty', 'quantity', 'no', 'count', 'unit', 'nos']):
                                quantity_columns.append(col)
                            # Look for item/description columns
                            elif any(keyword in col_str for keyword in ['item', 'description', 'work', 'activity', 'task', 'material', 'service', 'labour']):
                                item_columns.append(col)
                        
                        # If no cost columns found by keywords, look for numeric columns
                        if not cost_columns:
                            numeric_cols = df.select_dtypes(include=['number']).columns
                            # Check if numeric columns contain values that look like costs
                            for col in numeric_cols:
                                non_zero_values = df[col].dropna()
                                non_zero_values = non_zero_values[non_zero_values != 0]
                                if len(non_zero_values) > 0 and non_zero_values.max() > 10:  # Assume costs > 10
                                    cost_columns.append(col)
                        
                        text_content += f"IDENTIFIED COLUMNS:\n"
                        text_content += f"Cost/Amount columns: {cost_columns}\n"
                        text_content += f"Quantity columns: {quantity_columns}\n" 
                        text_content += f"Item/Description columns: {item_columns}\n\n"
                        
                        # Analyze cost data
                        sheet_total = 0
                        sheet_items = []
                        
                        # Process each cost column
                        for cost_col in cost_columns:
                            # Convert to numeric, replacing any text with NaN
                            df[cost_col] = pd.to_numeric(df[cost_col], errors='coerce')
                            
                            # Get all non-null, non-zero values
                            valid_costs = df[cost_col].dropna()
                            valid_costs = valid_costs[valid_costs != 0]  # Remove zeros
                            
                            if len(valid_costs) > 0:
                                col_total = valid_costs.sum()
                                col_max = valid_costs.max()
                                col_min = valid_costs.min()
                                col_avg = valid_costs.mean()
                                
                                text_content += f"COST ANALYSIS - {cost_col}:\n"
                                text_content += f"  Total: {col_total:,.2f}\n"
                                text_content += f"  Maximum: {col_max:,.2f}\n"
                                text_content += f"  Minimum: {col_min:,.2f}\n"
                                text_content += f"  Average: {col_avg:,.2f}\n"
                                text_content += f"  Valid entries: {len(valid_costs)} out of {len(df)}\n\n"
                                
                                sheet_total += col_total
                                
                                # Create cost items with descriptions
                                for idx, cost_value in valid_costs.items():
                                    if cost_value > 0:  # Only positive costs
                                        # Try to get item description
                                        item_desc = "Item"
                                        if len(item_columns) > 0:
                                            desc_value = df.loc[idx, item_columns[0]]
                                            if pd.notna(desc_value):
                                                item_desc = str(desc_value)[:60]  # Limit length
                                        
                                        # Add row number if no description
                                        if item_desc == "Item" or item_desc.strip() == "":
                                            item_desc = f"Row {idx + 1} Item"
                                        
                                        sheet_items.append((item_desc, cost_value, cost_col, sheet_name))
                        
                        # If no cost columns worked, try to find any numeric data
                        if sheet_total == 0:
                            text_content += "No cost columns identified by keywords. Checking all numeric columns...\n"
                            for col in df.columns:
                                try:
                                    numeric_data = pd.to_numeric(df[col], errors='coerce')
                                    valid_numbers = numeric_data.dropna()
                                    valid_numbers = valid_numbers[valid_numbers > 0]
                                    
                                    if len(valid_numbers) > 0 and valid_numbers.max() > 10:
                                        col_total = valid_numbers.sum()
                                        text_content += f"  Found numeric data in '{col}': {len(valid_numbers)} values, total: {col_total:,.2f}\n"
                                        
                                        # Add these as potential cost items
                                        for idx, value in valid_numbers.items():
                                            item_desc = f"Row {idx + 1} - {col}"
                                            if len(item_columns) > 0:
                                                desc_value = df.loc[idx, item_columns[0]]
                                                if pd.notna(desc_value):
                                                    item_desc = f"{str(desc_value)[:40]} ({col})"
                                            
                                            sheet_items.append((item_desc, value, col, sheet_name))
                                        
                                        sheet_total += col_total
                                except:
                                    continue
                        
                        # Show top cost items for this sheet
                        if sheet_items:
                            sheet_items_sorted = sorted(sheet_items, key=lambda x: x[1], reverse=True)
                            text_content += f"\nTOP 5 COST ITEMS IN {sheet_name}:\n"
                            for i, (item, cost, column, sheet) in enumerate(sheet_items_sorted[:5], 1):
                                text_content += f"  {i}. {item}: {cost:,.2f}\n"
                            text_content += "\n"
                        
                        # Add sample data
                        sample_rows = min(5, len(df))
                        text_content += f"SAMPLE DATA (first {sample_rows} rows):\n"
                        for index, row in df.head(sample_rows).iterrows():
                            row_text = " | ".join(str(value)[:30] if pd.notna(value) else "N/A" for value in row)
                            text_content += f"Row {index + 1}: {row_text}\n"
                        
                        if len(df) > sample_rows:
                            text_content += f"... and {len(df) - sample_rows} more rows\n"
                        
                        # Store sheet summary
                        sheet_summaries.append({
                            'name': sheet_name,
                            'total': sheet_total,
                            'rows': len(df),
                            'cost_columns': len(cost_columns) if cost_columns else 0,
                            'items_count': len(sheet_items)
                        })
                        
                        grand_total_budget += sheet_total
                        all_cost_items.extend(sheet_items)
                    
                    text_content += "\n" + "="*50 + "\n"
                
                # Overall summary
                text_content += f"\n=== OVERALL BOQ SUMMARY ===\n"
                text_content += f"CALCULATED TOTAL BUDGET: {grand_total_budget:,.2f}\n"
                text_content += f"Number of Sheets: {len(excel_data)}\n"
                text_content += f"Total Cost Items Found: {len(all_cost_items)}\n\n"
                
                # Top cost items across all sheets
                if all_cost_items:
                    all_cost_items.sort(key=lambda x: x[1], reverse=True)  # Sort by cost
                    text_content += "TOP 15 HIGHEST COST ITEMS (ACROSS ALL SHEETS):\n"
                    for i, (item, cost, column, sheet) in enumerate(all_cost_items[:15], 1):
                        text_content += f"{i}. {item}: {cost:,.2f} (Sheet: {sheet}, Column: {column})\n"
                    text_content += "\n"
                
                # Sheet-wise breakdown
                text_content += "SHEET-WISE BREAKDOWN:\n"
                for sheet_info in sheet_summaries:
                    percentage = (sheet_info['total'] / grand_total_budget * 100) if grand_total_budget > 0 else 0
                    text_content += f"• {sheet_info['name']}: {sheet_info['total']:,.2f} ({percentage:.1f}% of total, {sheet_info['items_count']} items)\n"
                
                # Cost distribution analysis
                if grand_total_budget > 0 and all_cost_items:
                    text_content += f"\nCOST DISTRIBUTION ANALYSIS:\n"
                    # Top 10 items percentage
                    top_10_total = sum(item[1] for item in all_cost_items[:10])
                    top_10_percentage = (top_10_total / grand_total_budget * 100)
                    text_content += f"Top 10 items represent: {top_10_percentage:.1f}% of total budget\n"
                    
                    # Average cost per item
                    avg_cost = grand_total_budget / len(all_cost_items)
                    text_content += f"Average cost per item: {avg_cost:,.2f}\n"
                
                logger.info(f"Extracted BOQ data: {len(text_content)} characters, Calculated Total Budget: {grand_total_budget:,.2f}")
                
            except Exception as e:
                logger.error(f"Failed to extract Excel content: {e}")
                raise HTTPException(status_code=500, detail=f"Failed to extract Excel content: {str(e)}")
                
        else:
            # Try to extract as plain text for other file types
            try:
                text_content = response.content.decode('utf-8', errors='ignore')
            except Exception as e:
                logger.error(f"Failed to extract text content: {e}")
                raise HTTPException(status_code=500, detail="Unsupported file type for text extraction")
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text content could be extracted from the document")
        
        # CLEAN AND FORMAT THE TEXT CONTENT
        def clean_text(text):
            """Clean and format extracted text for better AI processing"""
            # Replace multiple newlines with single newlines
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Replace multiple spaces with single spaces
            text = re.sub(r' {2,}', ' ', text)
            
            # Remove leading/trailing whitespace from each line
            lines = [line.strip() for line in text.split('\n')]
            
            # Remove empty lines
            lines = [line for line in lines if line]
            
            # Join lines back together with proper spacing
            cleaned_text = '\n'.join(lines)
            
            # Fix common formatting issues
            cleaned_text = cleaned_text.replace('\\n', '\n')
            cleaned_text = cleaned_text.replace('\\t', ' ')
            cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Replace multiple whitespace with single space
            cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)  # Proper paragraph separation
            
            return cleaned_text.strip()
        
        # Clean the extracted text
        text_content = clean_text(text_content)
        
        # Truncate content if too long (Gemini has token limits)
        max_chars = 50000
        if len(text_content) > max_chars:
            text_content = text_content[:max_chars] + "...\n[Content truncated due to length]"
            logger.info(f"Content truncated to {max_chars} characters")
        
        # Generate summary using Gemini AI
        if not GOOGLE_API_KEY:
            raise HTTPException(status_code=503, detail="AI summarization service is not configured")
        
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # Enhanced prompt for Excel/BOQ files
            if file_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"] or original_filename.lower().endswith(('.xls', '.xlsx')):
                prompt = f"""
                You are analyzing a BOQ (Bill of Quantities) or cost estimation document. Provide a comprehensive financial analysis and summary.
                
                **FORMATTING REQUIREMENTS:**
                - Use clear markdown formatting with proper headers
                - Use bullet points and tables where appropriate
                - Highlight important financial figures with **bold**
                
                **BOQ ANALYSIS STRUCTURE:**
                
                ## 📊 Executive Summary
                Provide a high-level overview of this BOQ/cost document.
                
                ## 💰 Financial Overview
                - **Total Project Budget:** [Extract the calculated total from the data]
                - **Number of Cost Categories:** [Count of different sections/sheets]
                - **Budget Range:** [Highest and lowest individual items]
                
                ## 🏆 Top Cost Items
                List the highest budget items found:
                - **Item 1:** Description and amount
                - **Item 2:** Description and amount
                - **Item 3:** Description and amount
                
                ## 📈 Budget Breakdown
                Analyze the cost distribution:
                - Major cost categories and their totals
                - Percentage breakdown if possible
                - Any patterns in the pricing structure
                
                ## 🔍 Key Insights
                - **Highest Cost Driver:** What category/item consumes most budget
                - **Budget Distribution:** How costs are spread across items
                - **Cost Patterns:** Any interesting observations about pricing
                - **Potential Risks:** Items with unusually high/low costs
                
                ## 📋 Document Structure
                - **Sheets/Sections:** What different parts exist
                - **Data Quality:** Completeness of cost information
                - **Organization:** How the BOQ is structured
                
                ## 💡 Recommendations
                Based on the cost analysis:
                - Budget optimization opportunities
                - Items requiring closer review
                - Cost management suggestions
                
                ---
                
                **Document Details:**
                - Name: {document_record.get('document_name', 'Untitled')}
                - Category: {document_record.get('document_category', 'Uncategorized')}
                - Type: BOQ/Cost Estimation Spreadsheet
                
                **Extracted BOQ Data:**
                {text_content}
                
                Please provide a detailed financial analysis following the structure above:
                """
            else:
                prompt = f"""
                Please provide a comprehensive and well-formatted summary of the following document. 
                
                **FORMATTING REQUIREMENTS:**
                - Use clear markdown formatting
                - Use proper headers (##, ###)
                - Use bullet points for lists
                - Use **bold** for important terms
                - Use line breaks for readability
                
                **CONTENT STRUCTURE:**
                
                ## Document Overview
                Brief description of what this document is about and its main purpose.
                
                ## Key Information
                - **Document Type:** [type]
                - **Period Covered:** [if applicable]
                - **Main Subject:** [subject]
                
                ## Main Points
                List the most important information, decisions, or findings:
                - Point 1
                - Point 2
                - Point 3
                
                ## Important Details
                Critical numbers, dates, names, specifications, or requirements:
                - **Detail 1:** Description
                - **Detail 2:** Description
                
                ## Action Items & Next Steps
                Any tasks, deadlines, responsibilities mentioned:
                - Action 1
                - Action 2
                
                ## Conclusion
                Main takeaways or outcomes
                
                ---
                
                **Document Information:**
                - Title: {document_record.get('document_name', 'Untitled')}
                - Category: {document_record.get('document_category', 'Uncategorized')}
                - File Type: {file_type}
                
                **Document Content:**
                {text_content}
                
                Please provide a clean, well-formatted summary following the structure above:
                """
            
            # Use sync version instead of async to avoid response object issues
            response = model.generate_content(prompt)
            summary = response.text.strip()
            
            # Additional cleaning of the AI response
            summary = summary.replace('\\n', '\n')
            summary = re.sub(r'\n{3,}', '\n\n', summary)
            
            logger.info(f"Generated formatted summary of {len(summary)} characters for document {doc_id}")
            
            return JSONResponse({
                "status": "success",
                "document_id": int(doc_id),
                "document_name": document_record.get("document_name", "Untitled"),
                "document_category": document_record.get("document_category", "Uncategorized"),
                "summary": summary,
                "content_length": len(text_content),
                "file_type": file_type,
                "version": latest_version.get("version", 1),
                "generated_at": datetime.now().isoformat(),
                "processing_info": {
                    "extraction_method": "PyMuPDF" if "pdf" in file_type.lower() else "python-docx" if "docx" in file_type.lower() else "pandas" if any(ext in file_type.lower() for ext in ["excel", "sheet"]) else "text",
                    "text_cleaned": True,
                    "formatting_applied": True,
                    "excel_sheets": len(excel_data) if 'excel_data' in locals() else None
                }
            })
            
        except Exception as e:
            logger.error(f"Gemini AI summarization failed: {e}")
            raise HTTPException(status_code=500, detail=f"AI summarization failed: {str(e)}")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in summarizeDocument: {e}")
        raise HTTPException(status_code=500, detail=f"Document summarization failed: {str(e)}")